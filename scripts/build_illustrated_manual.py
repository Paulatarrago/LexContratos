#!/usr/bin/env python3
from __future__ import annotations

import base64
import contextlib
import http.server
import json
import os
import random
import shutil
import socket
import socketserver
import struct
import subprocess
import tempfile
import threading
import time
import urllib.parse
import urllib.request
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "manual-ilustrado"
ASSET_DIR = OUT_DIR / "capturas"
DOCX_PATH = OUT_DIR / "Manual_LexContratos_Prueba_Piloto_Ilustrado.docx"
NOTE_PATH = OUT_DIR / "Nota_cambios_para_usuaria_piloto.docx"
PLAYWRIGHT_HEADLESS = Path("/private/tmp/pw-browsers/chromium_headless_shell-1217/chrome-headless-shell-mac-arm64/chrome-headless-shell")
CHROME = PLAYWRIGHT_HEADLESS if PLAYWRIGHT_HEADLESS.exists() else Path("/Applications/Google Chrome.app/Contents/MacOS/Google Chrome")


def find_free_port() -> int:
    with contextlib.closing(socket.socket(socket.AF_INET, socket.SOCK_STREAM)) as sock:
        sock.bind(("127.0.0.1", 0))
        return sock.getsockname()[1]


class LexHandler(http.server.SimpleHTTPRequestHandler):
    def __init__(self, *args, directory=None, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def log_message(self, *_):
        return

    def do_GET(self):
        parsed = urllib.parse.urlparse(self.path)
        path = parsed.path
        if path == "/lexconfig.local.js":
            body = b"window.lexSupabaseConfig = null;\n"
            self.send_response(200)
            self.send_header("Content-Type", "application/javascript; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
            return
        local_path = (ROOT / path.lstrip("/")).resolve()
        if path == "/" or not str(local_path).startswith(str(ROOT)) or not local_path.exists():
            self.path = "/index.html"
        super().do_GET()


def start_server(port: int):
    server = socketserver.TCPServer(("127.0.0.1", port), LexHandler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    return server


class CDP:
    def __init__(self, websocket_url: str):
        parsed = urllib.parse.urlparse(websocket_url)
        self.host = parsed.hostname or "127.0.0.1"
        self.port = parsed.port or 80
        self.path = parsed.path + (("?" + parsed.query) if parsed.query else "")
        self.sock = socket.create_connection((self.host, self.port), timeout=10)
        self._connect()
        self.next_id = 1

    def _connect(self):
        key = base64.b64encode(os.urandom(16)).decode("ascii")
        request = (
            f"GET {self.path} HTTP/1.1\r\n"
            f"Host: {self.host}:{self.port}\r\n"
            "Upgrade: websocket\r\n"
            "Connection: Upgrade\r\n"
            f"Sec-WebSocket-Key: {key}\r\n"
            "Sec-WebSocket-Version: 13\r\n\r\n"
        )
        self.sock.sendall(request.encode("ascii"))
        data = b""
        while b"\r\n\r\n" not in data:
            data += self.sock.recv(4096)
        if b" 101 " not in data.split(b"\r\n", 1)[0]:
            raise RuntimeError("Chrome DevTools no acepto la conexion WebSocket")

    def _send_frame(self, payload: bytes):
        header = bytearray([0x81])
        length = len(payload)
        if length < 126:
            header.append(0x80 | length)
        elif length < (1 << 16):
            header.append(0x80 | 126)
            header.extend(struct.pack("!H", length))
        else:
            header.append(0x80 | 127)
            header.extend(struct.pack("!Q", length))
        mask = os.urandom(4)
        header.extend(mask)
        masked = bytes(byte ^ mask[i % 4] for i, byte in enumerate(payload))
        self.sock.sendall(bytes(header) + masked)

    def _recv_exact(self, n: int) -> bytes:
        data = b""
        while len(data) < n:
            chunk = self.sock.recv(n - len(data))
            if not chunk:
                raise RuntimeError("Conexion WebSocket cerrada")
            data += chunk
        return data

    def _recv_frame(self) -> str | None:
        first, second = self._recv_exact(2)
        opcode = first & 0x0F
        masked = second & 0x80
        length = second & 0x7F
        if length == 126:
            length = struct.unpack("!H", self._recv_exact(2))[0]
        elif length == 127:
            length = struct.unpack("!Q", self._recv_exact(8))[0]
        mask = self._recv_exact(4) if masked else b""
        payload = self._recv_exact(length)
        if masked:
            payload = bytes(byte ^ mask[i % 4] for i, byte in enumerate(payload))
        if opcode == 8:
            raise RuntimeError("Chrome cerro la conexion")
        if opcode == 9:
            return None
        if opcode != 1:
            return None
        return payload.decode("utf-8")

    def call(self, method: str, params: dict | None = None, timeout: float = 10):
        msg_id = self.next_id
        self.next_id += 1
        self._send_frame(json.dumps({"id": msg_id, "method": method, "params": params or {}}).encode("utf-8"))
        deadline = time.time() + timeout
        while time.time() < deadline:
            raw = self._recv_frame()
            if not raw:
                continue
            message = json.loads(raw)
            if message.get("id") == msg_id:
                if "error" in message:
                    raise RuntimeError(message["error"])
                return message.get("result", {})
        raise TimeoutError(method)

    def eval(self, expression: str, await_promise: bool = False):
        return self.call(
            "Runtime.evaluate",
            {
                "expression": expression,
                "awaitPromise": await_promise,
                "returnByValue": True,
                "userGesture": True,
            },
        )

    def navigate(self, url: str, wait: float = 1.5):
        self.call("Page.navigate", {"url": url})
        time.sleep(wait)

    def screenshot(self, path: Path):
        result = self.call("Page.captureScreenshot", {"format": "png", "fromSurface": True, "captureBeyondViewport": False}, timeout=20)
        path.write_bytes(base64.b64decode(result["data"]))

    def close(self):
        with contextlib.suppress(Exception):
            self.sock.close()


ANNOTATION_JS = r"""
(() => {
  window.__lexClearAnnotations = () => {
    document.querySelectorAll('.lex-manual-annotation-layer').forEach((node) => node.remove());
  };
  window.__lexAnnotate = (items) => {
    window.__lexClearAnnotations();
    const layer = document.createElement('div');
    layer.className = 'lex-manual-annotation-layer';
    Object.assign(layer.style, {
      position: 'fixed',
      inset: '0',
      zIndex: '2147483647',
      pointerEvents: 'none',
      fontFamily: 'Arial, Helvetica, sans-serif'
    });
    const svg = document.createElementNS('http://www.w3.org/2000/svg', 'svg');
    svg.setAttribute('width', '100%');
    svg.setAttribute('height', '100%');
    svg.style.position = 'absolute';
    svg.style.inset = '0';
    svg.innerHTML = `
      <defs>
        <marker id="arrow" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth">
          <path d="M0,0 L0,6 L9,3 z" fill="#c77700"></path>
        </marker>
      </defs>`;
    layer.append(svg);
    document.body.append(layer);
    items.forEach((item, index) => {
      const target = document.querySelector(item.selector);
      if (!target) return;
      const rect = target.getBoundingClientRect();
      const targetX = item.targetX ?? (rect.left + rect.width / 2);
      const targetY = item.targetY ?? (rect.top + rect.height / 2);
      const x = item.x;
      const y = item.y;
      const width = item.width || 270;
      const label = document.createElement('div');
      label.textContent = `${index + 1}. ${item.text}`;
      Object.assign(label.style, {
        position: 'absolute',
        left: `${x}px`,
        top: `${y}px`,
        maxWidth: `${width}px`,
        padding: '10px 12px',
        borderRadius: '12px',
        background: '#fff8e6',
        border: '2px solid #c77700',
        color: '#13233a',
        boxShadow: '0 10px 26px rgba(9, 38, 77, .18)',
        fontSize: '15px',
        lineHeight: '1.25',
        fontWeight: '700'
      });
      layer.append(label);
      const labelX = x + (targetX < x ? 0 : width);
      const labelY = y + 22;
      const line = document.createElementNS('http://www.w3.org/2000/svg', 'line');
      line.setAttribute('x1', labelX);
      line.setAttribute('y1', labelY);
      line.setAttribute('x2', targetX);
      line.setAttribute('y2', targetY);
      line.setAttribute('stroke', '#c77700');
      line.setAttribute('stroke-width', '4');
      line.setAttribute('marker-end', 'url(#arrow)');
      line.setAttribute('stroke-linecap', 'round');
      svg.append(line);
      const ring = document.createElement('div');
      Object.assign(ring.style, {
        position: 'absolute',
        left: `${Math.max(0, rect.left - 6)}px`,
        top: `${Math.max(0, rect.top - 6)}px`,
        width: `${rect.width + 12}px`,
        height: `${rect.height + 12}px`,
        border: '3px solid #c77700',
        borderRadius: `${Math.min(18, Math.max(8, rect.height / 4))}px`,
        boxShadow: '0 0 0 9999px rgba(255, 255, 255, 0)',
      });
      layer.append(ring);
    });
  };
})();
"""


def wait_for(cdp: CDP, expression: str, timeout=8):
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            result = cdp.eval(expression)
            if result.get("result", {}).get("value"):
                return True
        except Exception:
            pass
        time.sleep(0.2)
    raise TimeoutError(expression)


def annotate_and_capture(cdp: CDP, filename: str, items: list[dict], wait=0.25) -> Path:
    cdp.eval(f"window.__lexAnnotate({json.dumps(items, ensure_ascii=False)});")
    time.sleep(wait)
    path = ASSET_DIR / filename
    cdp.screenshot(path)
    cdp.eval("window.__lexClearAnnotations();")
    return path


def wrapped_lines(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_callout(draw: ImageDraw.ImageDraw, number: int, text: str, box: tuple[int, int, int, int], target: tuple[int, int], font: ImageFont.ImageFont):
    x, y, w, h = box
    fill = "#fff8e6"
    border = "#c77700"
    text_color = "#13233a"
    draw.rounded_rectangle((x, y, x + w, y + h), radius=14, fill=fill, outline=border, width=3)
    lines = wrapped_lines(draw, f"{number}. {text}", font, w - 22)
    ty = y + 10
    for line in lines[:4]:
        draw.text((x + 11, ty), line, fill=text_color, font=font)
        ty += 19
    sx = x + w if target[0] > x + w else x
    sy = y + h // 2
    draw.line((sx, sy, target[0], target[1]), fill=border, width=4)
    draw.ellipse((target[0] - 9, target[1] - 9, target[0] + 9, target[1] + 9), outline=border, width=4)


def draw_signature_callouts(path: Path) -> Path:
    image = Image.open(path).convert("RGB")
    draw = ImageDraw.Draw(image)
    font_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    font = ImageFont.truetype(font_path, 16) if Path(font_path).exists() else ImageFont.load_default()
    callouts = [
        ("Nombre del firmante: persona física que firmará", (45, 275, 315, 68), (448, 437)),
        ("Correo: aquí llega la solicitud de Dropbox Sign", (45, 390, 315, 68), (670, 437)),
        ("Rol: representante legal u otro rol aplicable", (45, 505, 315, 68), (848, 437)),
        ("Agregar firmante si participa otra persona", (45, 620, 315, 68), (720, 586)),
        ("Estado: confirma si Dropbox Sign está conectado", (1120, 520, 310, 72), (590, 640)),
        ("Enviar prueba o enviar a firma según el modo configurado", (1120, 640, 310, 82), (720, 689)),
    ]
    for index, (text, box, target) in enumerate(callouts, 1):
        draw_callout(draw, index, text, box, target, font)
    image.save(path)
    return path


def launch_chrome(debug_port: int, user_dir: Path):
    cmd = [
        str(CHROME),
        "--headless=new",
        f"--remote-debugging-port={debug_port}",
        f"--user-data-dir={user_dir}",
        "--disable-gpu",
        "--no-sandbox",
        "--disable-dev-shm-usage",
        "--single-process",
        "--no-zygote",
        "--no-first-run",
        "--no-default-browser-check",
        "--window-size=1440,1000",
        "about:blank",
    ]
    return subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)


def connect_to_chrome(debug_port: int) -> CDP:
    list_url = f"http://127.0.0.1:{debug_port}/json/list"
    for _ in range(50):
        try:
            with urllib.request.urlopen(list_url, timeout=1) as response:
                targets = json.load(response)
            page = next((item for item in targets if item.get("type") == "page"), targets[0])
            return CDP(page["webSocketDebuggerUrl"])
        except Exception:
            time.sleep(0.2)
    raise RuntimeError("No se pudo conectar a Chrome")


def capture_screenshots() -> list[tuple[str, Path, list[str]]]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for old in ASSET_DIR.glob("*.png"):
        old.unlink()

    server_port = find_free_port()
    debug_port = find_free_port()
    server = start_server(server_port)
    chrome_dir = Path(tempfile.mkdtemp(prefix="lex-chrome-"))
    chrome = launch_chrome(debug_port, chrome_dir)
    cdp = None
    base = f"http://127.0.0.1:{server_port}"
    captures: list[tuple[str, Path, list[str]]] = []
    try:
      cdp = connect_to_chrome(debug_port)
      cdp.call("Page.enable")
      cdp.call("Runtime.enable")
      cdp.call("Emulation.setDeviceMetricsOverride", {"width": 1440, "height": 1000, "deviceScaleFactor": 1, "mobile": False})
      cdp.navigate(base + "/", 2)
      cdp.eval(ANNOTATION_JS)
      wait_for(cdp, "Boolean(document.querySelector('.landing-nav-actions button[data-route-link=\"/login\"]'))")
      captures.append((
        "1. Portada y acceso",
        annotate_and_capture(cdp, "01_portada_acceso.png", [
          {"selector": ".landing-nav-actions button[data-route-link='/login']", "text": "Entrar o registrarse: abre el acceso privado", "x": 1010, "y": 92},
          {"selector": ".landing-nav-actions a[href='#contacto']", "text": "Solicitar acceso: pide activación si aún no tienes cuenta", "x": 770, "y": 92},
          {"selector": ".auth-badge", "text": "Mensaje principal: extracción de datos desde Word, PDF, Excel, JPG y correos", "x": 810, "y": 210, "width": 360},
          {"selector": ".auth-manual", "text": "Manual rápido: resume el flujo de trabajo", "x": 820, "y": 560, "width": 310},
        ]),
        [
          "El botón Entrar o registrarse lleva al acceso privado.",
          "Solicitar acceso sirve para pedir activación o soporte.",
          "La portada explica qué hace LexContratos sin mostrar información privada.",
        ],
      ))

      cdp.navigate(base + "/login", 1.2)
      wait_for(cdp, "Boolean(document.querySelector('#login-form'))")
      cdp.eval(ANNOTATION_JS)
      captures.append((
        "2. Inicio de sesión",
        annotate_and_capture(cdp, "02_login.png", [
          {"selector": "#login-email", "text": "Correo de la cuenta registrada", "x": 880, "y": 405},
          {"selector": "#login-password", "text": "Contraseña personal", "x": 880, "y": 505},
          {"selector": "#remember-session", "text": "Mantener sesión iniciada en este equipo", "x": 850, "y": 600},
          {"selector": "#login-form .primary-action", "text": "Entrar: valida usuario, licencia y permisos", "x": 880, "y": 690},
          {"selector": "#show-register", "text": "Crear usuario: registro inicial para que administración active la licencia", "x": 820, "y": 800, "width": 330},
          {"selector": "#show-recover", "text": "Recuperar contraseña: envía enlace de restablecimiento", "x": 500, "y": 835, "width": 330},
        ]),
        [
          "La sesión requiere correo, contraseña y licencia activa.",
          "Crear usuario no activa la licencia automáticamente; administración la habilita.",
          "Mantener sesión iniciada evita firmarse cada vez en un equipo confiable.",
        ],
      ))

      cdp.eval("document.querySelector('#show-register').click();")
      time.sleep(0.4)
      captures.append((
        "3. Crear usuario",
        annotate_and_capture(cdp, "03_crear_usuario.png", [
          {"selector": "#register-name", "text": "Nombre de la persona usuaria", "x": 880, "y": 392},
          {"selector": "#register-email", "text": "Correo que aparecerá en el panel de administración", "x": 880, "y": 492},
          {"selector": "#register-password", "text": "Contraseña que usará para entrar", "x": 880, "y": 592},
          {"selector": "#register-form .primary-action", "text": "Crear usuario: deja la cuenta pendiente de licencia", "x": 880, "y": 700},
          {"selector": "#auth-register .show-login", "text": "Ya tengo cuenta: vuelve al inicio de sesión", "x": 870, "y": 795},
        ]),
        [
          "El registro crea una cuenta pendiente.",
          "La administradora recibe o consulta el registro para activar licencia.",
          "El usuario entra cuando su licencia queda activa.",
        ],
      ))

      cdp.navigate(base + "/", 0.6)
      cdp.eval("""
        localStorage.setItem('lexcontratos_session', JSON.stringify({ email: 'usuario.prueba@lexcontratos.local', date: new Date().toISOString() }));
        localStorage.setItem('lexcontratos_current_user', 'usuario.prueba@lexcontratos.local');
        localStorage.setItem('lexcontratos_last_login_email', 'usuario.prueba@lexcontratos.local');
      """)
      cdp.navigate(base + "/app", 2)
      cdp.eval(ANNOTATION_JS)
      wait_for(cdp, "!document.querySelector('#app-shell')?.classList.contains('is-hidden')")
      captures.append((
        "4. Área principal de trabajo",
        annotate_and_capture(cdp, "04_app_acciones_principales.png", [
          {"selector": "#open-user-guide", "text": "Manual de Uso: guía completa dentro de la app", "x": 920, "y": 24},
          {"selector": "#toggle-archive", "text": "Expedientes: abre el archivo de carpetas y contratos", "x": 45, "y": 245, "width": 290},
          {"selector": "#open-template-picker", "text": "Seleccionar formato: elige una plantilla precargada", "x": 315, "y": 245, "width": 300},
          {"selector": "label[for='template-import']", "text": "Importar draft propio: sube tu formato base", "x": 650, "y": 245, "width": 300},
          {"selector": "#new-template", "text": "Nuevo formato: redacta una plantilla desde cero", "x": 975, "y": 245, "width": 300},
        ]),
        [
          "La tira superior concentra las acciones para iniciar trabajo.",
          "Expedientes abre el archivo documental.",
          "Seleccionar o importar formato inicia el contrato de trabajo.",
        ],
      ))

      cdp.eval("document.querySelector('#open-template-picker').click();")
      time.sleep(0.5)
      captures.append((
        "5. Selector de formatos",
        annotate_and_capture(cdp, "05_selector_machotes.png", [
          {"selector": "#template-search", "text": "Buscar formato por tipo de contrato", "x": 910, "y": 195},
          {"selector": ".template-card", "text": "Tarjeta de formato: muestra nombre, categoría y campos", "x": 920, "y": 345, "width": 330},
          {"selector": ".template-card .use-template", "text": "Usar formato: crea una copia de trabajo", "x": 920, "y": 505, "width": 330},
        ]),
        [
          "El selector evita tener una lista enorme ocupando la pantalla.",
          "Al usar un formato, se trabaja una copia para no alterar la base.",
          "El buscador ayuda a encontrar rápido el contrato correcto.",
        ],
      ))

      cdp.eval("document.querySelector('.template-card')?.click();")
      time.sleep(0.8)
      wait_for(cdp, "Boolean(document.querySelector('#save-location-dialog')?.open)")
      captures.append((
        "6. Guardar como en expediente",
        annotate_and_capture(cdp, "06_guardar_como.png", [
          {"selector": "#save-location-file-name", "text": "Nombre del contrato de trabajo", "x": 735, "y": 142, "width": 310},
          {"selector": "#save-location-selected", "text": "Ruta seleccionada dentro del archivo", "x": 735, "y": 214, "width": 310},
          {"selector": "#save-location-browser", "text": "Explorador tipo Finder: entra a carpetas y subcarpetas", "x": 930, "y": 360, "width": 350},
          {"selector": "#save-location-confirm", "text": "Guardar aquí: crea el contrato en esa ubicación", "x": 980, "y": 885, "width": 320},
        ]),
        [
          "Guardar como define desde el inicio dónde vivirá el contrato.",
          "La ruta funciona como archivo de trabajo por cliente, proveedor, personal o documentos de partes.",
          "Dentro del explorador se pueden crear y organizar carpetas.",
        ],
      ))

      cdp.eval("""
        const name = document.querySelector('#save-location-file-name');
        if (name) {
          name.value = 'Contrato de prueba piloto';
          name.dispatchEvent(new Event('input', { bubbles: true }));
        }
        document.querySelector('#save-location-confirm')?.click();
      """)
      time.sleep(1.2)
      wait_for(cdp, "Boolean(document.querySelector('#contract-editor')?.value?.trim())")
      captures.append((
        "7. Editor y pasos obligatorios",
        annotate_and_capture(cdp, "07_editor_pasos.png", [
          {"selector": "#review-fields", "text": "1. Detectar campos editables: paso obligatorio", "x": 720, "y": 325, "width": 330},
          {"selector": "#fill-contract", "text": "2. Integrar manualmente: completa solo datos faltantes", "x": 715, "y": 410, "width": 340},
          {"selector": "#critical-review", "text": "3. Revisión crítica IA: observaciones finales", "x": 940, "y": 485, "width": 300},
          {"selector": "#export-word", "text": "Exportar Word: genera versión editable con formato legal", "x": 1070, "y": 575, "width": 330},
          {"selector": "#save-as-contract", "text": "Guardar como: cambia nombre o ubicación del contrato", "x": 920, "y": 735, "width": 340},
          {"selector": "#format-font", "text": "Fuente legal predeterminada: Georgia", "x": 40, "y": 665, "width": 300},
        ]),
        [
          "El flujo recomendado avanza de izquierda a derecha.",
          "Detectar campos editables es obligatorio antes de extraer o integrar datos.",
          "Guardar como está junto al documento porque afecta el archivo de trabajo.",
        ],
      ))

      cdp.eval("document.querySelector('#review-fields')?.click();")
      time.sleep(0.7)
      cdp.eval("document.querySelector('.party-upload-section')?.scrollIntoView({block: 'center'});")
      time.sleep(0.4)
      captures.append((
        "8. Carga de documentos por parte",
        annotate_and_capture(cdp, "08_documentos_por_parte.png", [
          {"selector": "#extract-data", "text": "Extraer e integrar datos: lee documentos y prellena campos", "x": 905, "y": 180, "width": 350},
          {"selector": ".role-drop", "text": "Caja de una parte: sube solo documentos de ese rol", "x": 70, "y": 515, "width": 330},
          {"selector": ".role-drop:nth-of-type(2)", "text": "Caja de la otra parte: evita mezclar información", "x": 940, "y": 515, "width": 330},
          {"selector": ".section-help", "text": "La ayuda explica qué documentos cargar y qué hará la extracción", "x": 820, "y": 285, "width": 390},
        ]),
        [
          "Los documentos deben separarse por parte contractual.",
          "La extracción intenta integrar lo que encuentre con soporte.",
          "Lo que falte queda para captura manual.",
        ],
      ))

      cdp.eval("window.scrollTo(0,0); document.querySelector('#fill-contract')?.click();")
      time.sleep(0.8)
      captures.append((
        "9. Datos faltantes manuales",
        annotate_and_capture(cdp, "09_datos_faltantes.png", [
          {"selector": "#assistant-pane", "text": "Panel lateral: muestra solo datos pendientes", "x": 645, "y": 155, "width": 330},
          {"selector": "#integrate-manual-data", "text": "Integrar datos capturados: inserta todos los campos llenos de una vez", "x": 565, "y": 300, "width": 360},
          {"selector": "#clear-generales", "text": "Borrar datos: limpia información integrada para cargar datos nuevos", "x": 560, "y": 380, "width": 350},
          {"selector": "#close-fields", "text": "Cerrar panel: también se oculta al dar clic fuera", "x": 1025, "y": 92, "width": 330},
        ]),
        [
          "Este panel no obliga a guardar campo por campo.",
          "Se capturan varios datos y se integran juntos.",
          "Los campos completos dejan de aparecer como pendientes.",
        ],
      ))

      cdp.eval("document.querySelector('#close-fields')?.click(); document.querySelector('#critical-review')?.click();")
      time.sleep(0.5)
      captures.append((
        "10. Revisión crítica IA",
        annotate_and_capture(cdp, "10_revision_critica.png", [
          {"selector": "#critical-observations", "text": "Solo observaciones: no modifica el contrato", "x": 760, "y": 170, "width": 330},
          {"selector": "#critical-suggest", "text": "Proponer ajustes: genera sugerencias para aceptar o descartar", "x": 780, "y": 245, "width": 360},
          {"selector": "#critical-congruence", "text": "Congruencia: compara contrato, anexos y documentos soporte", "x": 760, "y": 320, "width": 360},
          {"selector": "#critical-legal", "text": "Legislación aplicable: señala alertas normativas por país", "x": 760, "y": 395, "width": 360},
          {"selector": "#critical-country", "text": "País de referencia para la revisión legal", "x": 760, "y": 470, "width": 360},
          {"selector": "#critical-review-output", "text": "Área de resultados: hallazgos y cambios sugeridos", "x": 780, "y": 590, "width": 360},
          {"selector": "#critical-review-dialog .icon-button", "text": "Cerrar revisión: vuelve al contrato sin aplicar cambios", "x": 840, "y": 75, "width": 330},
        ]),
        [
          "La revisión crítica no sustituye el criterio del usuario.",
          "Las sugerencias se pueden aceptar o descartar una por una.",
          "La revisión legal aplicable no sustituye un dictamen ni certifica vigencia normativa.",
          "Antes de exportar, conviene revisar observaciones importantes.",
        ],
      ))

      cdp.eval("""
        document.querySelector('#critical-review-dialog')?.close();
        document.querySelector('#send-signature')?.click();
      """)
      time.sleep(0.6)
      wait_for(cdp, "Boolean(document.querySelector('#signature-dialog')?.open)")
      cdp.eval("""
        const rows = Array.from(document.querySelectorAll('[data-signer-row]'));
        const sampleSigners = [
          { name: 'Ana Sofía Martínez Ríos', email: 'ana.demo@lexcontratos.com', role: 'Representante legal del Cliente' },
          { name: 'Luis Eduardo Carranza Mora', email: 'luis.demo@lexcontratos.com', role: 'Representante legal del Prestador de servicios' },
        ];
        rows.forEach((row, index) => {
          const signer = sampleSigners[index] || sampleSigners[0];
          row.querySelector('[name="signerName"]').value = signer.name;
          row.querySelector('[name="signerEmail"]').value = signer.email;
          row.querySelector('[name="signerRole"]').value = signer.role;
          row.querySelector('[name="signerOrder"]').value = String(index + 1);
        });
        const copy = document.querySelector('#signature-dialog-copy');
        if (copy) copy.textContent = 'Dropbox Sign está conectado. El contrato se enviará por correo a los firmantes capturados. En modo prueba, el envío no es legalmente vinculante.';
        const status = document.querySelector('#signature-status-label');
        if (status) status.textContent = 'Dropbox Sign conectado · modo prueba';
        const submit = document.querySelector('#signature-submit');
        if (submit) submit.textContent = 'Enviar prueba de firma';
      """)
      time.sleep(0.3)
      signature_path = ASSET_DIR / "11_firma_electronica.png"
      cdp.screenshot(signature_path)
      draw_signature_callouts(signature_path)
      captures.append((
        "11. Firma electrónica con Dropbox Sign",
        signature_path,
        [
          "Para enviar a Dropbox Sign sí se necesita nombre, correo y rol de cada firmante.",
          "El contrato puede seguir incompleto para una prueba, pero los datos de firmantes no pueden estar vacíos.",
          "En modo prueba, el envío no es legalmente vinculante.",
          "Cuando el botón diga Enviar prueba de firma, el correo sirve para validar el flujo, no para firmar legalmente.",
          "Cuando el envío funciona, el firmante recibe un correo de Dropbox Sign para firmar desde su navegador.",
          "Si Dropbox Sign limita correos externos en modo prueba, usa correos del dominio autorizado o cambia a modo real cuando proceda.",
        ],
      ))

      cdp.eval("document.querySelector('#signature-dialog')?.close(); document.querySelector('#toggle-archive')?.click();")
      time.sleep(0.6)
      captures.append((
        "12. Expedientes y archivo",
        annotate_and_capture(cdp, "12_expedientes.png", [
          {"selector": "#archive-drawer", "text": "Archivo contractual: carpetas, contratos y documentos", "x": 405, "y": 140, "width": 340},
          {"selector": "#folder-root", "text": "Tipo de carpeta raíz: Clientes, Proveedores, Empresas del Grupo, Personales o Documentos de las partes", "x": 420, "y": 310, "width": 380},
          {"selector": "#folder-name", "text": "Crear carpeta/subcarpeta: escribe nombre y presiona Enter", "x": 420, "y": 405, "width": 360},
          {"selector": "#saved-contracts", "text": "Contenido de la carpeta: contratos guardados y acciones disponibles", "x": 420, "y": 640, "width": 360},
          {"selector": "#close-archive", "text": "Cerrar expedientes: oculta la barra lateral", "x": 355, "y": 65, "width": 310},
        ]),
        [
          "El archivo funciona como explorador interno.",
          "Las carpetas ayudan a ordenar por cliente, proveedor, asunto o documentos de las partes.",
          "La barra puede cerrarse para volver al contrato amplio.",
        ],
      ))
    finally:
      if cdp:
          cdp.close()
      chrome.terminate()
      with contextlib.suppress(Exception):
          chrome.wait(timeout=5)
      server.shutdown()
      shutil.rmtree(chrome_dir, ignore_errors=True)
    return captures


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, color="102033"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_note(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF4FF")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(9, 38, 77)
    p.add_run(f" {body}")


def add_steps_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Paso", True, "FFFFFF")
    set_cell_text(hdr[1], "Qué hacer", True, "FFFFFF")
    set_cell_shading(hdr[0], "09264D")
    set_cell_shading(hdr[1], "09264D")
    for step, detail in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], step, True, "09264D")
        set_cell_text(cells[1], detail)


def build_docx(captures: list[tuple[str, Path, list[str]]]):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(26)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 1"].font.color.rgb = RGBColor(9, 38, 77)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(21, 91, 184)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Manual ilustrado de uso\nLexContratos")
    r.bold = True
    r.font.size = Pt(27)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor(9, 38, 77)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subtitle.add_run("Prueba piloto interna | Guía visual paso a paso")
    rs.font.size = Pt(13)
    rs.font.color.rgb = RGBColor(96, 112, 137)

    add_note(
        doc,
        "Uso de este manual.",
        "Sigue las pantallas en orden. Las flechas señalan dónde está cada botón y para qué sirve. LexContratos asiste el llenado y organización contractual; la revisión final siempre queda bajo control del usuario.",
    )

    doc.add_heading("Cambios incorporados en esta versión", level=1)
    for change in [
        "La palabra machote se sustituyó por formato o draft, para que el lenguaje sea más claro para usuarias internas.",
        "El formato legal predeterminado cambió a fuente Georgia 12, con texto justificado y márgenes moderados para exportación a Word.",
        "Cada documento exportado puede incluir en el pie las iniciales del usuario que lo trabajó.",
        "Al seleccionar un formato precargado, se trabaja sobre una copia editable y se pide dónde guardar el documento.",
        "La app conserva datos existentes del contrato cuando ya son correctos y solo integra información faltante o validada.",
        "La revisión crítica IA ahora permite revisar observaciones, proponer ajustes granulares, comparar congruencia entre contrato/anexos/documentos y revisar legislación aplicable por país.",
        "La revisión de legislación aplicable señala alertas normativas y formalidades a validar; no sustituye dictamen jurídico ni consulta oficial vigente.",
        "Se agregó selección entre Notario Público y Corredor Público. Si es notario, el instrumento esperado es escritura pública; si es corredor, póliza o acta, según corresponda.",
        "El módulo de firma electrónica con Dropbox Sign permite preparar firmantes y enviar solicitudes cuando la integración está configurada.",
        "El guardado y archivo contractual se organizan con Guardar como, carpetas, expedientes y una vista más parecida a Finder.",
    ]:
        doc.add_paragraph(change, style="List Bullet")

    doc.add_heading("Flujo general", level=1)
    add_steps_table(
        doc,
        [
            ("1", "Entrar o crear usuario. La cuenta queda pendiente hasta que administración active la licencia."),
            ("2", "Seleccionar o importar formato/draft. La app trabaja sobre una copia para proteger la plantilla base."),
            ("3", "Guardar como. Elegir carpeta o expediente desde el inicio."),
            ("4", "Detectar campos editables. Este paso es obligatorio antes de cargar documentos o integrar datos."),
            ("5", "Cargar documentos por parte. No mezclar información de cliente, proveedor, prestador o contraparte."),
            ("6", "Extraer e integrar datos. Validar lo integrado y dejar pendiente lo que no esté claro."),
            ("7", "Completar datos faltantes. Integrar todos los datos capturados de una vez."),
            ("8", "Editar cláusulas y revisar con IA. Aceptar o descartar observaciones bajo criterio propio; revisar congruencia y legislación aplicable cuando convenga."),
            ("9", "Exportar a Word. Generar una versión editable para circular, imprimir o conservar en expediente."),
            ("10", "Preparar firma. Capturar firmantes, revisar correos y enviar por Dropbox Sign cuando la integración esté conectada."),
        ],
    )

    doc.add_page_break()
    doc.add_heading("Guía visual por pantalla", level=1)
    for idx, (heading, image_path, bullets) in enumerate(captures, 1):
        if idx > 1:
            doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading(heading, level=2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(7.15))
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Reglas de prueba para usuarias piloto", level=1)
    rules = [
        "Usar siempre lexcontratos.com para pruebas reales. Las vistas locales sirven solo para desarrollo.",
        "No subir documentos reales de terceros si no son necesarios para la prueba piloto.",
        "Separar documentos por parte. Si se mezclan, la extracción puede asignar datos al rol incorrecto.",
        "Presionar Detectar campos editables cada vez que se cambie el formato/draft o se borren campos del contrato.",
        "Revisar siempre los datos integrados antes de exportar a Word.",
        "Si se usa la revisión de legislación aplicable, tratarla como alerta de apoyo y no como dictamen definitivo.",
        "Si se usa corredor público, verificar que el documento sea póliza o acta, no escritura pública salvo que el usuario confirme otra cosa.",
        "Para enviar por Dropbox Sign, capturar nombre, correo y rol de cada firmante. No basta con tener el contrato abierto.",
        "En modo prueba de Dropbox Sign, el envío no es legalmente vinculante y puede estar limitado a correos del dominio autorizado.",
        "Reportar cualquier error con captura de pantalla, contrato utilizado, documentos cargados y una breve explicación.",
    ]
    for rule in rules:
        doc.add_paragraph(rule, style="List Bullet")

    add_note(
        doc,
        "Frase clave para soporte.",
        "Si algo no se entiende, pedir a la usuaria que indique en qué paso se atoró: acceso, formato/draft, guardar como, detectar campos, documentos por parte, extracción, datos faltantes, revisión IA, exportación o expedientes.",
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


def build_note_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(9, 38, 77)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Nota para prueba piloto de LexContratos")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor(9, 38, 77)

    doc.add_paragraph(
        "Hola. Te compartimos esta versión de LexContratos para prueba interna. La idea es validar si la herramienta realmente ayuda a reducir captura manual, evitar errores en generales y ordenar el trabajo contractual en expedientes."
    )
    doc.add_paragraph(
        "LexContratos no sustituye el criterio jurídico ni toma decisiones por el usuario. Ayuda a extraer datos, detectar campos pendientes, comparar información, organizar documentos y preparar una versión editable para revisión."
    )

    doc.add_heading("Qué cambió en esta versión", level=1)
    changes = [
        "El lenguaje cambió de machote a formato/draft.",
        "Al elegir un formato precargado, se genera una copia de trabajo y se pide dónde guardarla.",
        "El editor y la exportación a Word usan Georgia como fuente legal predeterminada.",
        "Los documentos exportados pueden llevar las iniciales del usuario en el pie.",
        "La extracción intenta integrar datos automáticamente y deja visibles solo los faltantes.",
        "Los datos manuales se pueden capturar en conjunto e integrar de una vez.",
        "Se reforzó la revisión crítica IA: observaciones, cambios granulares, congruencia entre contrato/anexos y revisión de legislación aplicable.",
        "Se agregó selección entre Notario Público y Corredor Público; si es corredor, se usa póliza o acta como instrumento.",
        "Se agregó el flujo de firma electrónica con Dropbox Sign: captura de firmantes, orden de firma y envío cuando la integración está conectada.",
        "El archivo de expedientes queda organizado por carpetas y subcarpetas para facilitar búsqueda.",
    ]
    for item in changes:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Qué revisar durante la prueba", level=1)
    checks = [
        "Si el flujo resulta claro desde seleccionar formato hasta exportar.",
        "Si Detectar campos editables se entiende como paso obligatorio.",
        "Si los documentos por parte evitan confusión de datos.",
        "Si la extracción identifica correctamente partes, representantes, fechas, poderes, RFC, domicilios y montos.",
        "Si los campos pendientes son útiles o sobran.",
        "Si la revisión crítica IA da observaciones accionables.",
        "Si la revisión de congruencia ayuda a encontrar diferencias entre contrato y anexos.",
        "Si el Word exportado se ve presentable para circular.",
        "Si el envío de firma llega correctamente al correo del firmante y el flujo de Dropbox Sign se entiende.",
    ]
    for item in checks:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Cómo reportar observaciones", level=1)
    doc.add_paragraph(
        "Cuando algo falle o no se entienda, por favor comparte: captura de pantalla, paso donde ocurrió, formato usado, tipo de documentos cargados y qué esperabas que hiciera la app. No es necesario mandar documentos sensibles si basta describir el caso."
    )

    add_note(
        doc,
        "Gracias por probarla.",
        "Tus comentarios son especialmente valiosos porque esta versión está pensada para validar el uso real antes de seguir endureciendo automatizaciones, firma digital y consulta normativa oficial.",
    )
    doc.save(NOTE_PATH)
    return NOTE_PATH


def main():
    captures = capture_screenshots()
    docx = build_docx(captures)
    note = build_note_docx()
    print(json.dumps({"docx": str(docx), "note": str(note), "screenshots": [str(path) for _, path, _ in captures]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
